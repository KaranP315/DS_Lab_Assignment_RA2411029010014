import os
from docx import Document
from docx.shared import Pt
import re

def create_doc():
    doc = Document()
    doc.add_heading('Data Science Lab Assignment', 0)
    
    # Details
    doc.add_paragraph("Name: Karan Pullagal")
    doc.add_paragraph("Roll Number: RA2411029010014")
    doc.add_paragraph("Class: B.Tech CSE (Core)")
    
    p_repo = doc.add_paragraph("GitHub Repository Link: ")
    # Clickable link placeholder
    p_repo.add_run("https://github.com/KaranPullagal/DS_Lab_Assignment_RA2411029010014").underline = True
    
    doc.add_paragraph("Question Numbers Attempted: Q1, Q2, Q6")
    doc.add_page_break()
    
    def add_question_section(q_num, problem, approach, code_file, output_file, challenges, insights):
        doc.add_heading(f"{q_num}", level=1)
        doc.add_heading("Problem Statement", level=2)
        doc.add_paragraph(problem)
        
        doc.add_heading("Approach / Logic used", level=2)
        doc.add_paragraph(approach)
        
        doc.add_heading("Complete Python Code", level=2)
        with open(code_file, 'r', encoding='utf-8') as f:
            code_text = f.read()
        p_code = doc.add_paragraph(code_text)
        p_code.style = 'No Spacing'
        for run in p_code.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            
        doc.add_heading("Output / Results", level=2)
        try:
            with open(output_file, 'r', encoding='utf-16') as f:
                out_text = f.read()
        except UnicodeError:
            with open(output_file, 'r', encoding='utf-8') as f:
                out_text = f.read()
        p_out = doc.add_paragraph(out_text)
        p_out.style = 'No Spacing'
        for run in p_out.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
        doc.add_heading("Challenges faced and how solved", level=2)
        doc.add_paragraph(challenges)
        
        doc.add_heading("Key insights / observations", level=2)
        doc.add_paragraph(insights)
        doc.add_page_break()

    # Q1
    q1_prob = ("Q1. Smart Missing Value Handling with Business Rules.\n"
               "Fill missing SalesAmount using median of ProductCategory and Region, or overall Category median. "
               "Create Imputation_Method column.")
    q1_app = "Grouped data by (ProductCategory, Region) for regional median and (ProductCategory) for category median. Used boolean masks to conditionally fill missing values and record imputation method."
    q1_chall = "Filtering logic to apply conditions properly without overriding values unnecessarily. Solved using dataframe `.loc` with boolean masking."
    q1_ins = "Regional median mapping handles values more explicitly by controlling geographical variance. Category medians act as a safety net when regional subsets are missing."
    
    add_question_section("Q1", q1_prob, q1_app, "Q1_missing_value_handling.py", "Q1_out.txt", q1_chall, q1_ins)

    # Q2
    q2_prob = ("Q2. Complex String Manipulation with Duplicate Handling.\n"
               "Remove titles, split into First, Middle, Last names, format into lower case firstinitial + lastname."
               " Append number to handle duplicate usernames.")
    q2_app = "Applied regex to extract and remove titles. Split the remaining text taking care of variable-length names using a function. Iteratively built username strings and tracked uniqueness using a dictionary counter."
    q2_chall = "Handling names with no middle names or varying name lengths. Addressed by counting string splits and conditionally allocating Name parts."
    q2_ins = "Using dictionary for counting username appearances operates in O(1) time complexity, efficiently solving duplication issues and keeping the generator robust for large logs."
    add_question_section("Q2", q2_prob, q2_app, "Q2_string_manipulation.py", "Q2_out.txt", q2_chall, q2_ins)
    
    # Q6
    q6_prob = ("Q6. Multi-File Data Wrangling.\n"
               "Merge student files with exam score files. Calculate average marks per Department per Subject. Pivot data.")
    q6_app = "Used `pd.merge` on StudentID. Then grouped by 'Department' and 'Subject' to find the mean marks. Re-structured using `DataFrame.pivot` with Department as index and Subject as columns."
    q6_chall = "Deciding on `how` inner vs outer join constraint. Inner join selected as only students with recorded marks needed computation."
    q6_ins = "Pivoting immediately turns high-volume dimension rows into interpretable aggregated cross-tabulations making it easier to analyze performance comparatively."
    add_question_section("Q6", q6_prob, q6_app, "Q6_multi_file_wrangling.py", "Q6_out.txt", q6_chall, q6_ins)
    
    doc.save("RA2411029010014_KaranPullagal_DS_Lab_Assignment.docx")

if __name__ == "__main__":
    create_doc()
